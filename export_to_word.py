import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys
import os
from docx.shared import RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Unified Style Utilities ---
def apply_table_style(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    return table

def apply_header_style(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
    set_cell_background(cell, '4F81BD')

# --- Override table creation to use unified style ---
def add_general_info_table(doc, headers, rows, border_color='4F81BD'):
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = 'Table Grid'
    set_table_borders(table, border_color)
    # Set first column width
    for row in table.rows:
        row.cells[0].width = Cm(4)
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].merge(hdr_cells[1])
    hdr_cells[0].text = headers[0] if headers[0].strip() else 'General Information'
    apply_header_style(hdr_cells[0])
    # Data rows
    for i, row in enumerate(rows):
        row_cells = table.rows[i+1].cells
        for j, cell in enumerate(row):
            row_cells[j].text = safe_value(cell)
            for paragraph in row_cells[j].paragraphs:
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    apply_table_style(table)
    return table

def add_table_with_headers(doc, headers, rows, header_bg='4F81BD', header_fg=RGBColor(255,255,255), border_color='4F81BD', no_headers=False):
    table = doc.add_table(rows=1 if not no_headers else 0, cols=len(headers))
    table.style = 'Table Grid'
    set_table_borders(table, border_color)
    if not no_headers:
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            apply_header_style(hdr_cells[i])
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = safe_value(cell)
            for paragraph in row_cells[i].paragraphs:
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_before = 0
                paragraph.paragraph_format.space_after = 0
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    apply_table_style(table)
    return table

# --- Unified section heading style ---
def add_numbered_section_heading(doc, number, text):
    heading = doc.add_paragraph()
    run = heading.add_run(f"{number}. {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    heading.style.font.color.rgb = RGBColor(0,0,0)
    heading.paragraph_format.line_spacing = 1
    heading.paragraph_format.space_before = 0
    heading.paragraph_format.space_after = 0
    return heading

def add_section_heading(doc, text, level=1):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    heading.style.font.color.rgb = RGBColor(0,0,0)
    heading.paragraph_format.line_spacing = 1
    heading.paragraph_format.space_before = 0
    heading.paragraph_format.space_after = 0
    return heading

def add_key_value_table(doc, data):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light List Accent 1'
    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value) if value is not None else ""
    doc.add_paragraph()  # Add space after table

def add_bullet_list(doc, items, title=None, style='List Bullet'):
    if title:
        doc.add_paragraph(title, style=style)
    for item in items:
        doc.add_paragraph(str(item), style=style)

def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_text_color(cell, rgb):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = rgb

def set_table_borders(table, color_hex):
    tbl = table._tbl
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}Border')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tbl.tblPr.append(border)

def set_paragraph_format(paragraph):
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        if run._element.rPr is not None:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(11)

def set_table_font_and_spacing(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph)

def safe_value(val):
    return str(val) if val not in [None, '', [], {}] else 'N/A'

def add_installation_method_table(doc, method, value):
    # Bolded subheading for method
    p = doc.add_paragraph()
    run = p.add_run(f"{method}:")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    set_paragraph_format(p)
    # Prepare table columns
    columns = ["Command", "Explanation", "Note", "When to use", "Platform"]
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    set_table_borders(table, '4F81BD')
    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255,255,255)
                run.font.name = 'Times New Roman'
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                run.font.size = Pt(11)
            set_paragraph_format(paragraph)
        set_cell_background(hdr_cells[i], '4F81BD')
    # Add rows
    if isinstance(value, list):
        for item in value:
            if not isinstance(item, dict):
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = safe_value(item.get('command', ''))
            row_cells[1].text = safe_value(item.get('explanation', ''))
            row_cells[2].text = safe_value(item.get('note', ''))
            row_cells[3].text = safe_value(item.get('when_to_use', ''))
            row_cells[4].text = safe_value(item.get('platform', ''))
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph)
    elif isinstance(value, dict):
        row_cells = table.add_row().cells
        row_cells[0].text = safe_value(value.get('command', ''))
        row_cells[1].text = safe_value(value.get('explanation', ''))
        row_cells[2].text = safe_value(value.get('note', ''))
        row_cells[3].text = safe_value(value.get('when_to_use', ''))
        row_cells[4].text = safe_value(value.get('platform', ''))
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph)
    else:
        # Fallback: just print value
        row_cells = table.add_row().cells
        row_cells[0].text = safe_value(value)
        for i in range(1, len(columns)):
            row_cells[i].text = ''
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                set_paragraph_format(paragraph)
    set_table_font_and_spacing(table)
    add_section_spacing(doc)

def add_installation_summary(doc, summary):
    if not summary or summary in ['N/A', None, '', [], {}]:
        doc.add_paragraph('N/A')
        return
    for method, value in summary.items():
        add_installation_method_table(doc, method, value)

def add_blue_left_col_table(doc, rows, left_col_width_cm=4.5, border_color='4F81BD', blue=True):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'
    set_table_borders(table, border_color)
    table.autofit = False
    # Calculate right column width for A4 with 1.27cm margins
    page_width_cm = 21.0
    margin_cm = 1.27
    right_col_width_cm = page_width_cm - 2 * margin_cm - left_col_width_cm
    for i, row in enumerate(rows):
        row_cells = table.rows[i].cells
        # Set left column (field name) with non-breaking spaces
        field_name = safe_value(row[0]).replace(' ', '\u00A0')
        row_cells[0].text = field_name
        row_cells[0].width = Cm(left_col_width_cm)
        row_cells[1].width = Cm(right_col_width_cm)
        for paragraph in row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                if blue:
                    run.font.color.rgb = RGBColor(255,255,255)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        if blue:
            set_cell_background(row_cells[0], '4F81BD')
        # Set right column (value)
        row_cells[1].text = safe_value(row[1])
        for paragraph in row_cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
    return table

def json_to_professional_word(json_path):
    import os
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    tool_name = data.get('Name')
    if not tool_name and 'General Info' in data:
        tool_name = data['General Info'].get('Name', 'output')
    if not tool_name:
        tool_name = 'output'
    output_dir = os.path.join('SampleOutputs', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    word_path = os.path.join(output_dir, f"{tool_name}.docx")

    doc = Document()
    # Set narrow margins (1.27 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    # Title (no square brackets)
    title = doc.add_paragraph()
    run = title.add_run("Tool Metadata Report (by MetadataFetcher)")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_format(title)

    # 1. General Information
    add_numbered_section_heading(doc, 1, 'General Information')
    gen_info = data.get('General Info', {})
    gen_rows = [
        ["Name", gen_info.get("Name", "N/A")],
        ["Use Case", gen_info.get("Use Case", "N/A")],
        ["Homepage", gen_info.get("Homepage", "N/A")],
        ["Description", gen_info.get("Description", "N/A")],
    ]
    add_blue_left_col_table(doc, gen_rows, left_col_width_cm=3.2, border_color='4F81BD', blue=True)
    add_section_spacing(doc)

    # 2. Documentation
    add_numbered_section_heading(doc, 2, 'Documentation')
    doc_info = data.get('Documentation', {})
    doc_rows = [
        ["Main Documentation", doc_info.get('Main Documentation', 'N/A')],
        ["Top Documentation Links", '\n'.join(doc_info.get('Top Links', [])) if doc_info.get('Top Links') else 'N/A']
    ]
    add_blue_left_col_table(doc, doc_rows, left_col_width_cm=5.2, border_color='4F81BD', blue=True)
    add_section_spacing(doc)

    # 3. Installation
    add_numbered_section_heading(doc, 3, 'Installation')
    install = data.get('Installation', {})
    install_links = install.get('Links', [])
    install_summary = install.get('Summary', {})
    install_links_str = '\n'.join(install_links) if install_links else 'N/A'
    # Installation Summary as a single cell (key: value, newline separated), show None for [] or None, and pretty-print commands
    def pretty_install_summary(summary):
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.shared import RGBColor
        import io
        lines = []
        for k, v in summary.items():
            if v is None or v == []:
                lines.append((k, None))
            elif isinstance(v, list):
                if not v:
                    lines.append((k, None))
                else:
                    for item in v:
                        if isinstance(item, dict):
                            # Bold method name
                            lines.append((k, 'BOLD'))
                            if 'command' in item:
                                lines.append((None, f"  Command: {item['command']}"))
                            if 'explanation' in item and item['explanation']:
                                lines.append((None, f"  Explanation: {item['explanation']}"))
                            if 'note' in item and item['note']:
                                lines.append((None, f"  Note: {item['note']}"))
                            if 'when_to_use' in item and item['when_to_use']:
                                lines.append((None, f"  When to use: {item['when_to_use']}"))
                            if 'platform' in item and item['platform']:
                                lines.append((None, f"  Platform: {item['platform']}"))
                            if 'more_info' in item and item['more_info']:
                                lines.append((None, f"  Link for Reference: {item['more_info']}"))
                        else:
                            lines.append((k, f"{item}"))
            elif isinstance(v, dict):
                lines.append((k, 'BOLD'))
                for subk, subv in v.items():
                    lines.append((None, f"  {subk}: {subv}"))
            else:
                lines.append((k, f"{v}"))
        # Now, build a string for the cell, but return a list of (label, value) for custom rendering
        return lines if lines else [(None, 'None')]

    if install_summary:
        install_summary_lines = pretty_install_summary(install_summary)
    else:
        install_summary_lines = [(None, 'N/A')]

    # Instead of install_summary_str, build the cell with runs for bold method names
    install_rows = [
        ["Installation Links", install_links_str],
        ["Installation Summary", None]  # We'll fill this cell below
    ]
    table = add_blue_left_col_table(doc, install_rows, left_col_width_cm=3.8, border_color='4F81BD', blue=False)
    # Fill the Installation Summary cell with formatted content
    summary_cell = table.rows[1].cells[1]
    summary_cell.text = ''  # Clear default
    for label, value in install_summary_lines:
        p = summary_cell.add_paragraph()
        if label and value == 'BOLD':
            run = p.add_run(f"{label}:")
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        elif value:
            run = p.add_run(value)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
    add_section_spacing(doc)

    # 4. Other Links
    add_numbered_section_heading(doc, 4, 'Other Links')
    other = data.get('Other Links', {})
    all_doc_links = other.get('All Documentation Links', [])
    if all_doc_links:
        p = doc.add_paragraph()
        run = p.add_run('All Documentation Links:')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        set_paragraph_format(p)
        for link in all_doc_links:
            p = doc.add_paragraph(link)
            set_paragraph_format(p)
    all_install_links = other.get('All Installation Links', [])
    if all_install_links:
        # Add a single blank line between the two subsections if both exist
        if all_doc_links:
            blank = doc.add_paragraph()
            set_paragraph_format(blank)
        p = doc.add_paragraph()
        run = p.add_run('All Installation Links:')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        set_paragraph_format(p)
        for link in all_install_links:
            p = doc.add_paragraph(link)
            set_paragraph_format(p)
    doc.save(word_path)
    print(f"Exported to {word_path}")

def export_product_docx(json_path):
    import os
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    tool_name = data.get('Name', 'output')
    tool_type = data.get('Type', 'software').lower()
    output_dir = os.path.join('SampleOutputs', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    word_path = os.path.join(output_dir, f"{tool_name}.docx")

    doc = Document()
    # Set narrow margins (1.27 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"{tool_name} Metadata Report (by MetadataFetcher)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_format(title)

    # 1. General Information
    add_numbered_section_heading(doc, 1, 'General Information')
    gen_rows = [
        ["Name", data.get("Name", "N/A")],
        ["Type", data.get("Type", "N/A")],
        ["Description", data.get("Description", "N/A")],
        ["Official Site", data.get("Official Site", "N/A")],
    ]
    add_general_info_table(doc, [' ', ' '], gen_rows, border_color='4F81BD')
    add_section_spacing(doc)

    # 2. Product Details
    add_numbered_section_heading(doc, 2, 'Product Details')
    details_rows = []
    if tool_type == 'hardware':
        details_rows.extend([
            ["Versions", data.get("Versions", "N/A")],
            ["Price Range", data.get("Price Range", "N/A")],
            ["Color Options", data.get("Color Options", "N/A")],
            ["Compatibility", data.get("Compatibility", "N/A")],
            ["Accessories", data.get("Accessories", "N/A")],
        ])
    else:  # software
        details_rows.extend([
            ["Versions", data.get("Versions", "N/A")],
            ["Compatibility", data.get("Compatibility", "N/A")],
        ])
    # Only show if at least one field is not N/A
    if any(val and val != "N/A" for _, val in details_rows):
        add_general_info_table(doc, [' ', ' '], details_rows, border_color='4F81BD')
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    # 3. Key Features
    add_numbered_section_heading(doc, 3, 'Key Features')
    features = data.get("Key Features", "N/A")
    if features and features != "N/A":
        for feat in features.split(','):
            p = doc.add_paragraph(feat.strip())
            set_paragraph_format(p)
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    # 4. Where to Buy (hardware only)
    if tool_type == 'hardware':
        add_numbered_section_heading(doc, 4, 'Where to Buy')
        where = data.get("Where to Buy", "N/A")
        if where and where != "N/A":
            for w in where.split(','):
                p = doc.add_paragraph(w.strip())
                set_paragraph_format(p)
        else:
            p = doc.add_paragraph("N/A")
            set_paragraph_format(p)
        add_section_spacing(doc)

    # 5. Support/Reviews
    add_numbered_section_heading(doc, 5, 'Support/Reviews')
    support = data.get("Support/Reviews", "N/A")
    if support and support != "N/A":
        for s in support.split(','):
            p = doc.add_paragraph(s.strip())
            set_paragraph_format(p)
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    doc.save(word_path)
    print(f"Exported to {word_path}")

def export_borderline_docx(json_path):
    import os
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    product = data.get('Product Info', {})
    technical = data.get('Technical Info', {})
    tool_name = product.get('Name', 'output')
    tool_type = product.get('Type', 'software').lower()
    output_dir = os.path.join('SampleOutputs', 'docs')
    os.makedirs(output_dir, exist_ok=True)
    word_path = os.path.join(output_dir, f"{tool_name}.docx")

    doc = Document()
    # Set narrow margins (1.27 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"{tool_name} Metadata Report (by MetadataFetcher)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_format(title)

    # 1. General Information (Product)
    add_numbered_section_heading(doc, 1, 'General Information')
    gen_rows = [
        ["Name", product.get("Name", "N/A")],
        ["Type", product.get("Type", "N/A")],
        ["Description", product.get("Description", "N/A")],
        ["Official Site", product.get("Official Site", "N/A")],
    ]
    add_general_info_table(doc, [' ', ' '], gen_rows, border_color='4F81BD')
    add_section_spacing(doc)

    # 2. Product Details
    add_numbered_section_heading(doc, 2, 'Product Details')
    details_rows = []
    if tool_type == 'hardware':
        details_rows.extend([
            ["Versions", product.get("Versions", "N/A")],
            ["Price Range", product.get("Price Range", "N/A")],
            ["Color Options", product.get("Color Options", "N/A")],
            ["Compatibility", product.get("Compatibility", "N/A")],
            ["Accessories", product.get("Accessories", "N/A")],
        ])
    else:  # software
        details_rows.extend([
            ["Versions", product.get("Versions", "N/A")],
            ["Compatibility", product.get("Compatibility", "N/A")],
        ])
    if any(val and val != "N/A" for _, val in details_rows):
        add_general_info_table(doc, [' ', ' '], details_rows, border_color='4F81BD')
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    # 3. Key Features
    add_numbered_section_heading(doc, 3, 'Key Features')
    features = product.get("Key Features", "N/A")
    if features and features != "N/A":
        for feat in features.split(','):
            p = doc.add_paragraph(feat.strip())
            set_paragraph_format(p)
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    # 4. Where to Buy (hardware only)
    if tool_type == 'hardware':
        add_numbered_section_heading(doc, 4, 'Where to Buy')
        where = product.get("Where to Buy", "N/A")
        if where and where != "N/A":
            for w in where.split(','):
                p = doc.add_paragraph(w.strip())
                set_paragraph_format(p)
        else:
            p = doc.add_paragraph("N/A")
            set_paragraph_format(p)
        add_section_spacing(doc)

    # 5. Support/Reviews
    add_numbered_section_heading(doc, 5, 'Support/Reviews')
    support = product.get("Support/Reviews", "N/A")
    if support and support != "N/A":
        for s in support.split(','):
            p = doc.add_paragraph(s.strip())
            set_paragraph_format(p)
    else:
        p = doc.add_paragraph("N/A")
        set_paragraph_format(p)
    add_section_spacing(doc)

    # 6. Technical Information (if available)
    if technical:
        add_numbered_section_heading(doc, 6, 'Technical Information')
        # General Info
        add_section_heading(doc, 'General Information', level=2)
        add_key_value_table(doc, technical.get('General Info', {}))
        add_section_spacing(doc)
        # Documentation
        add_section_heading(doc, 'Documentation', level=2)
        doc_info = technical.get('Documentation', {})
        if doc_info.get('Main Documentation'):
            doc.add_paragraph('Main Documentation:', style='List Bullet')
            doc.add_paragraph(doc_info['Main Documentation'], style='List Bullet 2')
        if doc_info.get('Top Links'):
            add_bullet_list(doc, doc_info['Top Links'], title='Top Documentation Links:', style='List Bullet 2')
        add_section_spacing(doc)
        # Installation (only if there are actual commands)
        add_section_heading(doc, 'Installation', level=2)
        install = technical.get('Installation', {})
        summary = install.get('Summary', {})
        has_commands = any(
            isinstance(v, (list, dict)) and v not in [[], {}, None, 'N/A']
            for v in summary.values()
        )
        if has_commands:
            add_installation_summary(doc, summary)
        else:
            p = doc.add_paragraph("N/A")
            set_paragraph_format(p)
        add_section_spacing(doc)
        # Other Links
        add_section_heading(doc, 'Other Links', level=2)
        other = technical.get('Other Links', {})
        if other.get('All Documentation Links'):
            add_bold_label(doc, 'All Documentation Links:')
            for link in other['All Documentation Links']:
                p = doc.add_paragraph(link)
                set_paragraph_format(p)
        if other.get('All Installation Links'):
            add_bold_label(doc, 'All Installation Links:')
            for link in other['All Installation Links']:
                p = doc.add_paragraph(link)
                set_paragraph_format(p)
        add_section_spacing(doc)

    doc.save(word_path)
    print(f"Exported to {word_path}")

def add_section_spacing(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    return p

def add_bold_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    return p

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python export_to_word.py <input_json>")
        sys.exit(1)
    json_path = sys.argv[1]
    if not os.path.exists(json_path):
        print(f"Input file {json_path} does not exist.")
        sys.exit(1)
    json_to_professional_word(json_path) 