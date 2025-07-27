import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class CreativeMediaReportGenerator:
    """Dedicated report generator for Creative and Media Tools"""
    
    def __init__(self):
        self.category = "Creative and Media Tools"
        self.output_dir = os.path.join('SampleOutputs', 'docs', 'creative_media')
        self.primary_color = '8E24AA'  # Purple for Creative Media
        self.secondary_color = 'BA68C8'  # Light purple
        
    def create_output_directory(self):
        """Create category-specific output directories"""
        os.makedirs(os.path.join(self.output_dir, 'docx'), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, 'pdf'), exist_ok=True)
        
    def set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
        
    def apply_creative_header_style(self, cell):
        """Apply Creative Media specific header styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        self.set_cell_background(cell, self.primary_color)
        
    def apply_creative_cell_style(self, cell):
        """Apply Creative Media specific cell styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            
    def add_creative_section_heading(self, doc, number, text):
        """Add Creative Media specific section heading"""
        heading = doc.add_paragraph()
        run = heading.add_run(f"{number}. {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        heading.paragraph_format.line_spacing = 1
        heading.paragraph_format.space_before = 0
        heading.paragraph_format.space_after = 0
        return heading
        
    def add_creative_info_table(self, doc, data):
        """Create Creative Media specific information table"""
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        
        # Set table to use full width within margins
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate full width: A4 width (21cm) - left margin (1.27cm) - right margin (1.27cm) = ~18.46cm
        # Distribute evenly: label column ~5.5cm, content column ~12.96cm for better balance
        for row in table.rows:
            row.cells[0].width = Cm(5.5)  # Label column (slightly smaller)
            row.cells[1].width = Cm(12.96)  # Content column (larger for detailed content)
            
        for i, (key, value) in enumerate(data):
            row_cells = table.rows[i].cells
            
            # Key cell (left column)
            row_cells[0].text = str(key)
            self.apply_creative_header_style(row_cells[0])
            
            # Value cell (right column)
            row_cells[1].text = str(value) if value not in [None, '', [], {}] else 'N/A'
            self.apply_creative_cell_style(row_cells[1])
            
        return table
        
    def add_creative_installation_section(self, doc, installation_data):
        """Add Creative Media specific installation section"""
        if not installation_data or not isinstance(installation_data, dict):
            doc.add_paragraph('N/A')
            return
            
        for method, commands in installation_data.items():
            if not commands:
                continue
                
            # Method heading
            method_heading = doc.add_paragraph()
            run = method_heading.add_run(f"Installation Method: {method}")
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(142, 36, 170)
            
            # Create table for commands
            if isinstance(commands, list):
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Command'
                header_cells[1].text = 'Explanation'
                header_cells[2].text = 'Note'
                
                for cell in header_cells:
                    self.apply_creative_header_style(cell)
                    
                # Add command rows
                for cmd in commands:
                    if isinstance(cmd, dict):
                        row_cells = table.add_row().cells
                        row_cells[0].text = cmd.get('command', '')
                        row_cells[1].text = cmd.get('explanation', '')
                        row_cells[2].text = cmd.get('note', '')
                        for cell in row_cells:
                            self.apply_creative_cell_style(cell)
                            
    def generate_creative_media_report(self, json_path):
        """Generate Creative Media specific report from JSON data following YAML template"""
        self.create_output_directory()
        
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Extract tool name
        general_info = data.get('General Info', {})
        tool_name = general_info.get('Name', 'creative_media_tool')
        
        # Create document
        doc = Document()
        
        # Set margins
        section = doc.sections[0]
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        
        # Title (no extra spacing)
        title = doc.add_paragraph()
        run = title.add_run(f"{tool_name} - Creative and Media Tool Report")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title.paragraph_format.space_after = Pt(0)  # Remove extra spacing after title
        
        # Subtitle (no extra spacing)
        subtitle = doc.add_paragraph()
        run = subtitle.add_run("Generated by MetadataFetcher - Creative Media Category")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        subtitle.paragraph_format.space_after = Pt(0)  # Remove extra spacing
        
        # Follow Creative_and_Media_Tools field template from tool_descriptions.yaml
        creative_media_fields = [
            ('Tool Name', general_info.get('Name')),
            ('Overview / Description', general_info.get('Description')),
            ('Main Functions', general_info.get('Use Case')),
            ('Supported File Formats', data.get('FileFormats', 'N/A')),
            ('System Requirements', data.get('SystemRequirements', 'N/A')),
            ('Installation & Setup', self._format_installation_methods(data.get('Installation', {}).get('Summary', {}))),
            ('Key Features', data.get('Features', 'N/A')),
            ('Plugin/Extension Support', data.get('Plugins', 'N/A')),
            ('User Interface Overview', data.get('UI', 'N/A')),
            ('Community & Tutorials', data.get('Community', 'N/A')),
            ('Licensing', data.get('License', 'N/A')),
            ('Latest Version / Release Date', data.get('Version', 'N/A')),
            ('Example Projects / Outputs', data.get('Examples', 'N/A')),
            ('References (Official Website, Docs, etc.)', self._format_references(data)),
            ('Other Supporting Links (Github, etc.)', self._format_supporting_links(data))
        ]
        
        # Create comprehensive table with all Creative Media fields (no section heading)
        self.add_creative_info_table(doc, creative_media_fields)
        
        # Save document
        output_path = os.path.join(self.output_dir, 'docx', f"{tool_name}.docx")
        doc.save(output_path)
        print(f"Creative Media report generated: {output_path}")
        return output_path
        
    def _format_installation_methods(self, installation_data):
        """Format installation methods for display"""
        if not installation_data:
            return 'N/A'
        
        methods = []
        for method, commands in installation_data.items():
            if commands:
                methods.append(f"{method}: {len(commands)} command(s)")
        
        return '; '.join(methods) if methods else 'N/A'
        
    def _format_references(self, data):
        """Format references for display"""
        references = []
        
        if data.get('General Info', {}).get('Homepage'):
            references.append(f"Official Website: {data['General Info']['Homepage']}")
        
        doc_links = data.get('Documentation', {}).get('Top Links', [])
        if doc_links:
            references.extend(doc_links[:2])
        
        return '\n'.join(references) if references else 'N/A'
        
    def _format_supporting_links(self, data):
        """Format supporting links (GitHub, etc.) for display"""
        supporting_links = []
        
        # Check for GitHub links
        other_links = data.get('Other Links', {})
        github_links = [link for link in other_links.get('All Documentation Links', []) 
                       if 'github.com' in link.lower()]
        
        if github_links:
            supporting_links.extend(github_links[:3])
        
        # Add other supporting links
        all_links = other_links.get('All Documentation Links', [])
        non_github_links = [link for link in all_links if 'github.com' not in link.lower()]
        if non_github_links:
            supporting_links.extend(non_github_links[:2])
        
        return '\n'.join(supporting_links) if supporting_links else 'N/A'

def export_creative_media_tool(json_path):
    """Export function for Creative Media tools"""
    generator = CreativeMediaReportGenerator()
    return generator.generate_creative_media_report(json_path)