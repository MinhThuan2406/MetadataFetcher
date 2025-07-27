import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DataScienceReportGenerator:
    """Dedicated report generator for Data Science and Analytics Tools"""
    
    def __init__(self):
        self.category = "Data Science and Analytics Tools"
        self.output_dir = os.path.join('SampleOutputs', 'docs', 'data_science')
        self.primary_color = '2E7D32'  # Green for Data Science
        self.secondary_color = '66BB6A'  # Light green
        
    def create_output_directory(self):
        """Create category-specific output directories"""
        os.makedirs(os.path.join(self.output_dir, 'docx'), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, 'pdf'), exist_ok=True)
        
    def set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
        
    def apply_data_science_header_style(self, cell):
        """Apply Data Science specific header styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        self.set_cell_background(cell, self.primary_color)
        
    def apply_data_science_cell_style(self, cell):
        """Apply Data Science specific cell styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            
    def add_data_science_section_heading(self, doc, number, text):
        """Add Data Science specific section heading"""
        heading = doc.add_paragraph()
        run = heading.add_run(f"{number}. {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        heading.paragraph_format.line_spacing = 1
        heading.paragraph_format.space_before = 0
        heading.paragraph_format.space_after = 0
        return heading
        
    def add_analytics_metrics_table(self, doc, data):
        """Create Data Science specific metrics table"""
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        
        # Set table to use full width within margins
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate full width: A4 width (21cm) - left margin (1.27cm) - right margin (1.27cm) = ~18.46cm
        # Distribute evenly: label column ~5.5cm, content column ~12.96cm for better balance
        for row in table.rows:
            row.cells[0].width = Cm(5.5)  # Label column (slightly smaller)
            row.cells[1].width = Cm(12.96)  # Content column (larger for detailed content)
            
        for i, (key, value) in enumerate(data):
            row_cells = table.rows[i].cells
            
            # Metric name (left column)
            row_cells[0].text = str(key)
            self.apply_data_science_header_style(row_cells[0])
            
            # Metric value (right column)
            row_cells[1].text = str(value) if value not in [None, '', [], {}] else 'N/A'
            self.apply_data_science_cell_style(row_cells[1])
            
        return table
        
    def add_package_installation_guide(self, doc, installation_data):
        """Add Data Science specific installation guide"""
        if not installation_data or not isinstance(installation_data, dict):
            doc.add_paragraph('No installation data available')
            return
            
        for method, commands in installation_data.items():
            if not commands:
                continue
                
            # Installation method heading
            p = doc.add_paragraph()
            run = p.add_run(f"📦 {method.replace('_', ' ').title()} Installation:")
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(102, 187, 106)  # Data Science secondary color
            
            # Create installation table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Headers with data science focus
            headers = ['Command', 'Purpose', 'Environment', 'Notes']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                self.apply_data_science_header_style(hdr_cells[i])
                
            # Add installation commands
            if isinstance(commands, list):
                for cmd in commands:
                    if isinstance(cmd, dict):
                        row_cells = table.add_row().cells
                        row_cells[0].text = cmd.get('command', '')
                        row_cells[1].text = cmd.get('explanation', 'Package installation')
                        row_cells[2].text = cmd.get('platform', 'Cross-platform')
                        row_cells[3].text = cmd.get('note', '')
                        for cell in row_cells:
                            self.apply_data_science_cell_style(cell)
                            
    def add_data_workflow_section(self, doc, links):
        """Add data workflow and analysis links section"""
        if not links:
            doc.add_paragraph('No workflow resources available')
            return
            
        # Add workflow links as bullet points
        for link in links[:5]:  # Limit to top 5 for readability
            p = doc.add_paragraph()
            run = p.add_run(f"• {link}")
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(46, 125, 50)
            p.paragraph_format.left_indent = Cm(0.5)
            
    def generate_data_science_report(self, json_path):
        """Generate Data Science specific report from JSON data following YAML template"""
        self.create_output_directory()
        
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Extract tool name
        general_info = data.get('General Info', {})
        tool_name = general_info.get('Name', 'data_science_tool')
        
        # Create document
        doc = Document()
        
        # Set margins for data-heavy content
        section = doc.sections[0]
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        
        # Title with data science theme (no extra spacing)
        title = doc.add_paragraph()
        run = title.add_run(f"{tool_name} - Data Science & Analytics Report")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title.paragraph_format.space_after = Pt(0)  # Remove extra spacing after title
        
        # Subtitle (no extra spacing)
        subtitle = doc.add_paragraph()
        run = subtitle.add_run("MetadataFetcher - Data Science Category Analysis")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        subtitle.paragraph_format.space_after = Pt(0)  # Remove extra spacing
        
        # Follow Data_Science_and_Analytics_Tools field template from tool_descriptions.yaml
        data_science_fields = [
            ('Tool Name', general_info.get('Name')),
            ('Overview / Description', general_info.get('Description')),
            ('Typical Applications', general_info.get('Use Case')),
            ('Data Formats Supported', data.get('DataFormats', 'N/A')),
            ('Visualization Capabilities', data.get('Visualization', 'N/A')),
            ('Integration with Other Libraries', data.get('Integrations', 'N/A')),
            ('Installation & Setup', self._format_installation_methods(data.get('Installation', {}).get('Summary', {}))),
            ('Key Features', data.get('Features', 'N/A')),
            ('Community & Ecosystem', data.get('Community', 'N/A')),
            ('Documentation & Learning Resources', self._format_documentation_links(data.get('Documentation', {}))),
            ('Licensing', data.get('License', 'N/A')),
            ('Latest Version / Release Date', data.get('Version', 'N/A')),
            ('Example Use Cases', data.get('Examples', 'N/A')),
            ('References (Official Website, Docs, etc.)', self._format_references(data)),
            ('Other Supporting Links (Github, etc.)', self._format_supporting_links(data))
        ]
        
        # Create comprehensive table with all Data Science fields (no section heading)
        self.add_analytics_metrics_table(doc, data_science_fields)
        
        # Save document
        output_path = os.path.join(self.output_dir, 'docx', f"{tool_name}.docx")
        doc.save(output_path)
        print(f"Data Science report generated: {output_path}")
        return output_path
        
    def _format_installation_methods(self, installation_data):
        """Format installation methods for display"""
        if not installation_data:
            return 'N/A'
        
        methods = []
        for method, commands in installation_data.items():
            if commands:
                methods.append(f"{method}: {len(commands)} command(s)")
        
        return '; '.join(methods) if methods else 'N/A'
        
    def _format_documentation_links(self, doc_data):
        """Format documentation links for display"""
        if not doc_data:
            return 'N/A'
        
        links = []
        if doc_data.get('Main Documentation'):
            links.append(f"Main: {doc_data['Main Documentation']}")
        
        top_links = doc_data.get('Top Links', [])
        if top_links:
            links.extend(top_links[:3])
        
        return '\n'.join(links) if links else 'N/A'
        
    def _format_references(self, data):
        """Format references for display"""
        references = []
        
        if data.get('General Info', {}).get('Homepage'):
            references.append(f"Official Website: {data['General Info']['Homepage']}")
        
        doc_links = data.get('Documentation', {}).get('Top Links', [])
        if doc_links:
            references.extend(doc_links[:2])
        
        return '\n'.join(references) if references else 'N/A'
        
    def _format_supporting_links(self, data):
        """Format supporting links (GitHub, etc.) for display"""
        supporting_links = []
        
        # Check for GitHub links
        other_links = data.get('Other Links', {})
        github_links = [link for link in other_links.get('All Documentation Links', []) 
                       if 'github.com' in link.lower()]
        
        if github_links:
            supporting_links.extend(github_links[:3])
        
        # Add other supporting links
        all_links = other_links.get('All Documentation Links', [])
        non_github_links = [link for link in all_links if 'github.com' not in link.lower()]
        if non_github_links:
            supporting_links.extend(non_github_links[:2])
        
        return '\n'.join(supporting_links) if supporting_links else 'N/A'

def export_data_science_tool(json_path):
    """Export function for Data Science tools"""
    generator = DataScienceReportGenerator()
    return generator.generate_data_science_report(json_path)