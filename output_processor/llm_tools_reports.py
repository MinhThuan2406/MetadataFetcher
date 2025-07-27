import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class LLMToolsReportGenerator:
    """Dedicated report generator for Large Language Models (LLM) Tools"""
    
    def __init__(self):
        self.category = "Large Language Models (LLM) Tools"
        self.output_dir = os.path.join('SampleOutputs', 'docs', 'llm_tools')
        self.primary_color = 'FF6F00'  # Orange for LLM
        self.secondary_color = 'FFB74D'  # Light orange
        
    def create_output_directory(self):
        """Create category-specific output directories"""
        os.makedirs(os.path.join(self.output_dir, 'docx'), exist_ok=True)
        os.makedirs(os.path.join(self.output_dir, 'pdf'), exist_ok=True)
        
    def set_cell_background(self, cell, color):
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
        
    def apply_llm_header_style(self, cell):
        """Apply LLM Tools specific header styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        self.set_cell_background(cell, self.primary_color)
        
    def apply_llm_cell_style(self, cell):
        """Apply LLM Tools specific cell styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            
    def add_llm_section_heading(self, doc, number, text):
        """Add LLM Tools specific section heading"""
        heading = doc.add_paragraph()
        run = heading.add_run(f"{number}. {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        heading.paragraph_format.line_spacing = 1
        heading.paragraph_format.space_before = 0
        heading.paragraph_format.space_after = 0
        return heading
        
    def add_llm_info_table(self, doc, data):
        """Create LLM Tools specific information table"""
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        
        # Set table to use full width within margins
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate full width: A4 width (21cm) - left margin (1.27cm) - right margin (1.27cm) = ~18.46cm
        # Distribute evenly: label column ~5.5cm, content column ~12.96cm for better balance
        for row in table.rows:
            row.cells[0].width = Cm(5.5)  # Label column (slightly smaller)
            row.cells[1].width = Cm(12.96)  # Content column (larger for detailed content)
            
        for i, (key, value) in enumerate(data):
            row_cells = table.rows[i].cells
            
            # Parameter name (left column)
            row_cells[0].text = str(key)
            self.apply_llm_header_style(row_cells[0])
            
            # Parameter value (right column)
            row_cells[1].text = str(value) if value not in [None, '', [], {}] else 'N/A'
            self.apply_llm_cell_style(row_cells[1])
            
        return table
        
    def add_model_setup_section(self, doc, installation_data):
        """Add LLM Tools specific model setup section"""
        if not installation_data or not isinstance(installation_data, dict):
            doc.add_paragraph('No model setup information available')
            return
            
        for method, commands in installation_data.items():
            if not commands:
                continue
                
            # Setup method heading
            p = doc.add_paragraph()
            run = p.add_run(f"🤖 {method.replace('_', ' ').title()} Setup:")
            run.bold = True
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 183, 77)  # LLM secondary color
            
            # Create setup table
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Headers
            headers = ['Setup Command', 'Purpose', 'Model Requirements', 'Usage Context']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                self.apply_llm_header_style(hdr_cells[i])
                
            # Add setup commands
            if isinstance(commands, list):
                for cmd in commands:
                    if isinstance(cmd, dict):
                        row_cells = table.add_row().cells
                        row_cells[0].text = cmd.get('command', '')
                        row_cells[1].text = cmd.get('explanation', 'Model installation')
                        row_cells[2].text = cmd.get('note', 'Standard requirements')
                        row_cells[3].text = cmd.get('when_to_use', 'General use')
                        for cell in row_cells:
                            self.apply_llm_cell_style(cell)
                            
    def add_model_capabilities_section(self, doc, links):
        """Add model capabilities and examples section"""
        if not links:
            doc.add_paragraph('No model capability documentation available')
            return
            
        # Add capability links with LLM context
        for i, link in enumerate(links[:5]):  # Limit to top 5 for focus
            p = doc.add_paragraph()
            run = p.add_run(f"🔗 Model Resource {i+1}: {link}")
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 111, 0)
            p.paragraph_format.left_indent = Cm(0.5)
            
    def generate_llm_tools_report(self, json_path):
        """Generate LLM Tools specific report from JSON data following YAML template"""
        self.create_output_directory()
        
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        # Extract tool name
        general_info = data.get('General Info', {})
        tool_name = general_info.get('Name', 'llm_tool')
        
        # Create document
        doc = Document()
        
        # Set margins
        section = doc.sections[0]
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        
        # Title (no extra spacing)
        title = doc.add_paragraph()
        run = title.add_run(f"{tool_name} - Large Language Models (LLM) Tool Report")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title.paragraph_format.space_after = Pt(0)  # Remove extra spacing after title
        
        # Subtitle (no extra spacing)
        subtitle = doc.add_paragraph()
        run = subtitle.add_run("Generated by MetadataFetcher - LLM Tools Category")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 128, 128)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        subtitle.paragraph_format.space_after = Pt(0)  # Remove extra spacing
        
        # Follow LLM_Tools field template from tool_descriptions.yaml
        llm_tools_fields = [
            ('Tool Name', general_info.get('Name')),
            ('Overview / Description', general_info.get('Description')),
            ('Supported Model Types', data.get('ModelTypes', 'N/A')),
            ('Key Features', data.get('Features', 'N/A')),
            ('Installation & Setup', self._format_installation_methods(data.get('Installation', {}).get('Summary', {}))),
            ('Integration with Other Tools/Frameworks', data.get('Integrations', 'N/A')),
            ('Model Deployment Options', data.get('Deployment', 'N/A')),
            ('API/SDK Availability', data.get('API', 'N/A')),
            ('Documentation & Tutorials', self._format_documentation_links(data.get('Documentation', {}))),
            ('Community & Support', data.get('Community', 'N/A')),
            ('Licensing', data.get('License', 'N/A')),
            ('Latest Version / Release Date', data.get('Version', 'N/A')),
            ('Example Use Cases / Demos', data.get('Examples', 'N/A')),
            ('References (Official Website, Docs, etc.)', self._format_references(data)),
            ('Other Supporting Links (Github, etc.)', self._format_supporting_links(data))
        ]
        
        # Create comprehensive table with all LLM Tools fields (no section heading)
        self.add_llm_info_table(doc, llm_tools_fields)
        
        # Save document
        output_path = os.path.join(self.output_dir, 'docx', f"{tool_name}.docx")
        doc.save(output_path)
        print(f"LLM Tools report generated: {output_path}")
        return output_path
        
    def _format_installation_methods(self, installation_data):
        """Format installation methods for display"""
        if not installation_data:
            return 'N/A'
        
        methods = []
        for method, commands in installation_data.items():
            if commands:
                methods.append(f"{method}: {len(commands)} command(s)")
        
        return '; '.join(methods) if methods else 'N/A'
        
    def _format_documentation_links(self, doc_data):
        """Format documentation links for display"""
        if not doc_data:
            return 'N/A'
        
        links = []
        if doc_data.get('Main Documentation'):
            links.append(f"Main: {doc_data['Main Documentation']}")
        
        top_links = doc_data.get('Top Links', [])
        if top_links:
            links.extend(top_links[:3])
        
        return '\n'.join(links) if links else 'N/A'
        
    def _format_references(self, data):
        """Format references for display"""
        references = []
        
        if data.get('General Info', {}).get('Homepage'):
            references.append(f"Official Website: {data['General Info']['Homepage']}")
        
        doc_links = data.get('Documentation', {}).get('Top Links', [])
        if doc_links:
            references.extend(doc_links[:2])
        
        return '\n'.join(references) if references else 'N/A'
        
    def _format_supporting_links(self, data):
        """Format supporting links (GitHub, etc.) for display"""
        supporting_links = []
        
        # Check for GitHub links
        other_links = data.get('Other Links', {})
        github_links = [link for link in other_links.get('All Documentation Links', []) 
                       if 'github.com' in link.lower()]
        
        if github_links:
            supporting_links.extend(github_links[:3])
        
        # Add other supporting links
        all_links = other_links.get('All Documentation Links', [])
        non_github_links = [link for link in all_links if 'github.com' not in link.lower()]
        if non_github_links:
            supporting_links.extend(non_github_links[:2])
        
        return '\n'.join(supporting_links) if supporting_links else 'N/A'

def export_llm_tool(json_path):
    """Export function for LLM Tools"""
    generator = LLMToolsReportGenerator()
    return generator.generate_llm_tools_report(json_path)