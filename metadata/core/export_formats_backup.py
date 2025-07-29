"""
Advanced output processor for UnifiedMetadata.
Exports metadata to DOCX, PDF, and TXT formats with professional formatting.
"""

import json
import yaml
from typing import Union
from .schema import UnifiedMetadata
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def load_schema():
    """Load the schema configuration."""
    try:
        with open('schema.yaml', 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        # Fallback schema if file not found
        return {
            'basic_info': ['tool_name', 'overview_description', 'key_features', 'use_cases', 'supported_platforms'],
            'technical_specs': ['dependencies', 'system_requirements', 'performance_considerations', 'supported_languages', 'supported_file_formats'],
            'setup_info': ['installation_setup', 'configuration_guide', 'quick_start_tutorial'],
            'documentation': ['official_documentation', 'tutorials_examples', 'community_tutorials', 'api_reference'],
            'community': ['community_support', 'forums_channels', 'ecosystem_packages'],
            'legal_versioning': ['licensing', 'latest_version', 'release_date', 'version_history'],
            'references': ['official_website', 'github_repository', 'documentation_links', 'additional_resources'],
            'category_fields': {},
            'field_descriptions': {},
            'field_types': {}
        }


def get_all_fields():
    """Get all fields from schema in the correct order as defined in schema.yaml."""
    schema = load_schema()
    
    # Get all fields from organized sections
    all_fields = []
    
    # Add fields from each section in order
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    for section in sections:
        section_fields = schema.get(section, [])
        all_fields.extend(section_fields)
    
    # Add category-specific fields
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        for field in fields:
            if field not in all_fields:
                all_fields.append(field)
    
    return all_fields


def get_field_display_name(field_name):
    """Get a proper display name for a field."""
    field_mapping = {
        'tool_name': 'Tool Name',
        'overview_description': 'Overview Description',
        'key_features': 'Key Features',
        'use_cases': 'Use Cases',
        'supported_platforms': 'Supported Platforms',
        'dependencies': 'Dependencies',
        'system_requirements': 'System Requirements',
        'performance_considerations': 'Performance Considerations',
        'supported_languages': 'Supported Languages',
        'supported_file_formats': 'Supported File Formats',
        'installation_setup': 'Installation Setup',
        'configuration_guide': 'Configuration Guide',
        'quick_start_tutorial': 'Quick Start Tutorial',
        'official_documentation': 'Official Documentation',
        'tutorials_examples': 'Tutorials and Examples',
        'community_tutorials': 'Community Tutorials',
        'api_reference': 'API Reference',
        'community_support': 'Community Support',
        'forums_channels': 'Forums and Channels',
        'ecosystem_packages': 'Ecosystem Packages',
        'licensing': 'Licensing',
        'latest_version': 'Latest Version',
        'release_date': 'Release Date',
        'version_history': 'Version History',
        'official_website': 'Official Website',
        'github_repository': 'GitHub Repository',
        'documentation_links': 'Documentation Links',
        'additional_resources': 'Additional Resources'
    }
    
    return field_mapping.get(field_name, field_name.replace('_', ' ').title())


def get_section_display_name(section_name):
    """Get a proper display name for a section."""
    section_mapping = {
        'basic_info': 'Basic Information',
        'technical_specs': 'Technical Specifications',
        'setup_info': 'Setup and Installation',
        'documentation': 'Documentation and Learning',
        'community': 'Community and Support',
        'legal_versioning': 'Legal and Versioning',
        'references': 'References and Links'
    }
    
    return section_mapping.get(section_name, section_name.replace('_', ' ').title())


def format_field_value(value):
    """Format a field value for display."""
    if isinstance(value, list):
        return ', '.join(str(item) for item in value)
    elif isinstance(value, dict):
        return json.dumps(value, indent=2)
    else:
        return str(value) if value is not None else "N/A"


def save_to_docx(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a DOCX file with professional formatting.
    Follows user preferences: numbered headers, 11pt Times New Roman, narrow margins, proper spacing.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create document
    doc = Document()
    
    # Set narrow margins (1.27 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    
    # Set up custom styles
    # Title style - 11pt Times New Roman, black, bold
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(11)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Heading 1 style - 11pt Times New Roman, black, bold, no spacing after
    heading_style = doc.styles['Heading 1']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(11)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading_style.paragraph_format.space_after = Pt(0)  # No spacing after
    heading_style.paragraph_format.space_before = Pt(0)  # No extra spacing
    
    # Normal style for body text - 11pt Times New Roman, readable spacing
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)  # Minimal spacing
    normal_style.paragraph_format.space_before = Pt(0)  # No extra spacing
    normal_style.paragraph_format.line_spacing = 1.2  # Readable spacing
    
    # Add title
    tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
    title = doc.add_heading(f"{tool_name} - Metadata Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process fields by sections with proper headers and logical ordering
    schema = load_schema()
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    section_number = 1
    
    # Process each section in order
    for section in sections:
        section_fields = schema.get(section, [])
        section_data = data.get(section, {})
        
        # Check if section has any meaningful content
        has_content = False
        for field in section_fields:
            value = section_data.get(field)
            if value and value != "" and value != [] and value != {}:
                has_content = True
                break
        
        if has_content:
            # Add section header
            section_title = get_section_display_name(section)
            doc.add_heading(f"{section_number}. {section_title}", level=1)
            section_number += 1
            
            # Process fields within this section
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    # Add field header (no extra spacing)
                    field_title = get_field_display_name(field)
                    doc.add_heading(f"  {field_title}", level=2)
                    
                    # Add field content in prose format (immediately after header)
                    if isinstance(value, list):
                        # Convert list to prose
                        if len(value) == 1:
                            doc.add_paragraph(str(value[0]))
                        else:
                            prose_text = f"{value[0]}"
                            for item in value[1:-1]:
                                prose_text += f", {item}"
                            if len(value) > 1:
                                prose_text += f", and {value[-1]}"
                            doc.add_paragraph(prose_text)
                    else:
                        formatted_value = format_field_value(value)
                        if formatted_value and formatted_value.strip():
                            doc.add_paragraph(formatted_value)
    
    # Process category-specific fields (at root level)
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        for field in fields:
            value = data.get(field)
            if value and value != "" and value != [] and value != {}:
                # Add field header (no extra spacing)
                field_title = get_field_display_name(field)
                doc.add_heading(f"{section_number}. {field_title}", level=1)
                
                # Add field content in prose format (immediately after header)
                if isinstance(value, list):
                    # Convert list to prose
                    if len(value) == 1:
                        doc.add_paragraph(str(value[0]))
                    else:
                        prose_text = f"{value[0]}"
                        for item in value[1:-1]:
                            prose_text += f", {item}"
                        if len(value) > 1:
                            prose_text += f", and {value[-1]}"
                        doc.add_paragraph(prose_text)
                else:
                    formatted_value = format_field_value(value)
                    if formatted_value and formatted_value.strip():
                        doc.add_paragraph(formatted_value)
                
                section_number += 1
    
    # Save document
    doc.save(file_path)


def save_to_pdf(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a PDF file with professional formatting.
    Follows user preferences: numbered headers, 11pt Times New Roman, narrow margins, proper spacing.
    """
    if not PDF_AVAILABLE:
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create PDF document with narrow margins (1.27 cm)
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=A4,
        leftMargin=1.27*cm,
        rightMargin=1.27*cm,
        topMargin=1.27*cm,
        bottomMargin=1.27*cm
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles matching user preferences
    # Title style - 11pt Times New Roman, bold, centered
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=12,
        alignment=1,  # Center
        textColor=colors.black
    )
    
    # Normal style for body text - 11pt Times New Roman, readable spacing
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        spaceAfter=6,
        spaceBefore=0,
        leading=14,  # Readable spacing
        textColor=colors.black,
        keepWithNext=False,
        keepWithPrevious=False
    )
    
    # Header style - 11pt Times New Roman, bold, no spacing after
    header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=0,
        spaceBefore=0,
        textColor=colors.black,
        keepWithNext=True,
        keepWithPrevious=False
    )
    
    # Title
    tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
    story.append(Paragraph(f"{tool_name} - Metadata Report", title_style))
    
    # Process fields by sections with proper headers and logical ordering
    schema = load_schema()
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    section_number = 1
    
    # Process each section in order
    for section in sections:
        section_fields = schema.get(section, [])
        section_data = data.get(section, {})
        
        # Check if section has any meaningful content
        has_content = False
        for field in section_fields:
            value = section_data.get(field)
            if value and value != "" and value != [] and value != {}:
                has_content = True
                break
        
        if has_content:
            # Add section header
            section_title = get_section_display_name(section)
            story.append(Paragraph(f"{section_number}. {section_title}", header_style))
            section_number += 1
            
            # Process fields within this section
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    # Add field header (no extra spacing)
                    field_title = get_field_display_name(field)
                    story.append(Paragraph(f"  {field_title}", header_style))
                    
                    # Add field content in prose format (immediately after header)
                    if isinstance(value, list):
                        # Convert list to prose
                        if len(value) == 1:
                            story.append(Paragraph(str(value[0]), normal_style))
                        else:
                            prose_text = f"{value[0]}"
                            for item in value[1:-1]:
                                prose_text += f", {item}"
                            if len(value) > 1:
                                prose_text += f", and {value[-1]}"
                            story.append(Paragraph(prose_text, normal_style))
                    else:
                        formatted_value = format_field_value(value)
                        if formatted_value and formatted_value.strip():
                            story.append(Paragraph(formatted_value, normal_style))
    
    # Process category-specific fields (at root level)
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        for field in fields:
            value = data.get(field)
            if value and value != "" and value != [] and value != {}:
                # Add field header (no extra spacing)
                field_title = get_field_display_name(field)
                story.append(Paragraph(f"{section_number}. {field_title}", header_style))
                
                # Add field content in prose format (immediately after header)
                if isinstance(value, list):
                    # Convert list to prose
                    if len(value) == 1:
                        story.append(Paragraph(str(value[0]), normal_style))
                    else:
                        prose_text = f"{value[0]}"
                        for item in value[1:-1]:
                            prose_text += f", {item}"
                        if len(value) > 1:
                            prose_text += f", and {value[-1]}"
                        story.append(Paragraph(prose_text, normal_style))
                else:
                    formatted_value = format_field_value(value)
                    if formatted_value and formatted_value.strip():
                        story.append(Paragraph(formatted_value, normal_style))
                
                section_number += 1
    
    # Build PDF
    doc.build(story)


def save_to_txt(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a TXT file with detailed, structured formatting.
    Follows the format shown in sample_report.txt with proper section headers and comprehensive content.
    """
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    with open(file_path, 'w', encoding='utf-8') as f:
        # Title
        tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
        f.write(f"{tool_name} - Metadata Report\n\n")
        
        # Process fields by sections with proper headers and logical ordering
        schema = load_schema()
        sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
        
        section_number = 1
        
        # Process each section in order
        for section in sections:
            section_fields = schema.get(section, [])
            section_data = data.get(section, {})
            
            # Check if section has any meaningful content
            has_content = False
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    has_content = True
                    break
            
            if has_content:
                # Add section header
                section_title = get_section_display_name(section)
                f.write(f"{section_number}. {section_title}\n")
                section_number += 1
                
                # Process fields within this section
                for field in section_fields:
                    value = section_data.get(field)
                    if value and value != "" and value != [] and value != {}:
                        # Add field header (no extra spacing)
                        field_title = get_field_display_name(field)
                        f.write(f"  {field_title}\n")
                        
                        # Add field content in prose format (immediately after header)
                        if isinstance(value, list):
                            # Convert list to prose
                            if len(value) == 1:
                                f.write(f"    {value[0]}\n")
                            else:
                                prose_text = f"{value[0]}"
                                for item in value[1:-1]:
                                    prose_text += f", {item}"
                                if len(value) > 1:
                                    prose_text += f", and {value[-1]}"
                                f.write(f"    {prose_text}\n")
                        else:
                            formatted_value = format_field_value(value)
                            if formatted_value and formatted_value.strip():
                                f.write(f"    {formatted_value}\n")
                        f.write("\n")
        
        # Process category-specific fields (at root level)
        category_fields = schema.get('category_fields', {})
        for category, fields in category_fields.items():
            for field in fields:
                value = data.get(field)
                if value and value != "" and value != [] and value != {}:
                    # Add field header (no extra spacing)
                    field_title = get_field_display_name(field)
                    f.write(f"{section_number}. {field_title}\n")
                    
                    # Add field content in prose format (immediately after header)
                    if isinstance(value, list):
                        # Convert list to prose
                        if len(value) == 1:
                            f.write(f"  {value[0]}\n")
                        else:
                            prose_text = f"{value[0]}"
                            for item in value[1:-1]:
                                prose_text += f", {item}"
                            if len(value) > 1:
                                prose_text += f", and {value[-1]}"
                            f.write(f"  {prose_text}\n")
                    else:
                        formatted_value = format_field_value(value)
                        if formatted_value and formatted_value.strip():
                            f.write(f"  {formatted_value}\n")
                    
                    section_number += 1
                    f.write("\n")


def get_available_formats() -> list:
    """Get list of available export formats."""
    formats = ['json', 'yaml', 'txt']
    
    if DOCX_AVAILABLE:
        formats.append('docx')
    
    if PDF_AVAILABLE:
        formats.append('pdf')
    
    return formats 