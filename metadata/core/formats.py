"""
Advanced output processor for UnifiedMetadata.
Exports metadata to DOCX and PDF formats with professional formatting.
"""

import json
from typing import Union
from .schema import UnifiedMetadata
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def save_to_docx(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a DOCX file with professional formatting.
    Follows user preferences: 11pt Times New Roman headers, prose format, organized content.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create document
    doc = Document()
    
    # Set up custom styles matching user preferences
    # Title style - 11pt Times New Roman, black, bold
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(11)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Heading 1 style - 11pt Times New Roman, black, bold, no extra spacing
    heading_style = doc.styles['Heading 1']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(11)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading_style.paragraph_format.space_after = Pt(0)  # No extra spacing
    
    # Normal style for body text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(10)
    
    # Add title
    title = doc.add_heading(f"{data.get('display_name', data.get('name', 'Tool'))} - Metadata Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add generation info in prose format
    doc.add_paragraph(f"This metadata report was generated on {datetime.now().strftime('%Y-%m-%d at %H:%M:%S')}. "
                     f"The information was compiled from the following sources: {', '.join(data.get('sources', ['online databases']))}.")
    
    # Basic Information - prose format
    doc.add_heading("Basic Information", level=1)
    
    # Present information in prose rather than bullet points
    basic_info = []
    if data.get('name'):
        basic_info.append(f"The tool is named '{data['name']}'")
    if data.get('display_name'):
        basic_info.append(f"with a display name of '{data['display_name']}'")
    if data.get('category'):
        basic_info.append(f"and is classified in the {data['category']} category")
    if data.get('version'):
        basic_info.append(f"with version {data['version']}")
    if data.get('latest_version'):
        basic_info.append(f"and the latest available version is {data['latest_version']}")
    
    if basic_info:
        doc.add_paragraph('. '.join(basic_info) + '.')
    
    # Description in prose
    if data.get('description'):
        doc.add_paragraph(f"The tool's description states: {data['description']}")
    
    # License and author information
    if data.get('license') or data.get('author') or data.get('maintainer'):
        doc.add_paragraph("Regarding licensing and authorship:")
        if data.get('license'):
            doc.add_paragraph(f"The tool is licensed under {data['license']}.")
        if data.get('author'):
            doc.add_paragraph(f"The author of this tool is {data['author']}.")
        if data.get('maintainer'):
            doc.add_paragraph(f"The current maintainer is {data['maintainer']}.")
    
    # Links section with proper labeling
    if data.get('links'):
        doc.add_heading("Links", level=1)
        doc.add_paragraph("The following links provide additional information about this tool:")
        for link in data['links']:
            link_title = link.get('title', 'Link')
            link_url = link.get('url', '')
            if link_title.lower() == 'documentation':
                doc.add_paragraph(f"Documentation: {link_url}")
            elif link_title.lower() == 'homepage':
                doc.add_paragraph(f"Homepage: {link_url}")
            elif link_title.lower() == 'repository':
                doc.add_paragraph(f"Repository: {link_url}")
            else:
                doc.add_paragraph(f"{link_title}: {link_url}")
    
    # Installation Methods with detailed instructions
    if data.get('installation_methods'):
        doc.add_heading("Installation Methods", level=1)
        doc.add_paragraph("The following installation methods are available for this tool:")
        for method in data['installation_methods']:
            method_name = method.get('method', 'Unknown')
            command = method.get('command', '')
            description = method.get('description', '')
            
            doc.add_paragraph(f"For {method_name} installation, use the following command: {command}")
            if description:
                doc.add_paragraph(f"Additional details: {description}")
    
    # Dependencies in prose format
    if data.get('dependencies'):
        doc.add_heading("Dependencies", level=1)
        doc.add_paragraph("This tool has the following dependencies:")
        for version, deps in data['dependencies'].items():
            doc.add_paragraph(f"For version {version}, the required dependencies are: {', '.join(deps)}")
    
    # Category-specific fields in prose
    if data.get('category_fields'):
        doc.add_heading("Additional Information", level=1)
        doc.add_paragraph("The following additional information is available for this tool:")
        for key, value in data['category_fields'].items():
            if value:
                field_name = key.replace('_', ' ').title()
                doc.add_paragraph(f"{field_name}: {value}")
    
    # Save document
    doc.save(file_path)


def save_to_pdf(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a PDF file with professional formatting.
    Follows user preferences: prose format, organized content, proper link labeling.
    """
    if not PDF_AVAILABLE:
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create PDF document
    doc = SimpleDocTemplate(file_path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles matching user preferences
    # Title style - 11pt Times New Roman, black, bold
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName='Times-Roman',
        fontSize=11,
        spaceAfter=20,
        alignment=1,  # Center
        textColor=colors.black
    )
    
    # Heading style - 11pt Times New Roman, black, bold, no extra spacing
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=6,
        spaceBefore=12,
        textColor=colors.black
    )
    
    # Normal style for body text
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10,
        spaceAfter=6,
        textColor=colors.black
    )
    
    # Title
    title = Paragraph(f"{data.get('display_name', data.get('name', 'Tool'))} - Metadata Report", title_style)
    story.append(title)
    
    # Generation info in prose format
    gen_info = f"This metadata report was generated on {datetime.now().strftime('%Y-%m-%d at %H:%M:%S')}. "
    gen_info += f"The information was compiled from the following sources: {', '.join(data.get('sources', ['online databases']))}."
    story.append(Paragraph(gen_info, normal_style))
    story.append(Spacer(1, 20))
    
    # Basic Information - prose format
    story.append(Paragraph("Basic Information", heading_style))
    
    # Present information in prose rather than tables
    basic_info = []
    if data.get('name'):
        basic_info.append(f"The tool is named '{data['name']}'")
    if data.get('display_name'):
        basic_info.append(f"with a display name of '{data['display_name']}'")
    if data.get('category'):
        basic_info.append(f"and is classified in the {data['category']} category")
    if data.get('version'):
        basic_info.append(f"with version {data['version']}")
    if data.get('latest_version'):
        basic_info.append(f"and the latest available version is {data['latest_version']}")
    
    if basic_info:
        story.append(Paragraph('. '.join(basic_info) + '.', normal_style))
    
    # Description in prose
    if data.get('description'):
        desc_text = f"The tool's description states: {data['description']}"
        if len(desc_text) > 500:
            desc_text = desc_text[:500] + "..."
        story.append(Paragraph(desc_text, normal_style))
    
    # License and author information
    if data.get('license') or data.get('author') or data.get('maintainer'):
        story.append(Paragraph("Regarding licensing and authorship:", normal_style))
        if data.get('license'):
            story.append(Paragraph(f"The tool is licensed under {data['license']}.", normal_style))
        if data.get('author'):
            story.append(Paragraph(f"The author of this tool is {data['author']}.", normal_style))
        if data.get('maintainer'):
            story.append(Paragraph(f"The current maintainer is {data['maintainer']}.", normal_style))
    
    story.append(Spacer(1, 15))
    
    # Links section with proper labeling
    if data.get('links'):
        story.append(Paragraph("Links", heading_style))
        story.append(Paragraph("The following links provide additional information about this tool:", normal_style))
        for link in data['links']:
            link_title = link.get('title', 'Link')
            link_url = link.get('url', '')
            if link_title.lower() == 'documentation':
                link_text = f"Documentation: {link_url}"
            elif link_title.lower() == 'homepage':
                link_text = f"Homepage: {link_url}"
            elif link_title.lower() == 'repository':
                link_text = f"Repository: {link_url}"
            else:
                link_text = f"{link_title}: {link_url}"
            
            if len(link_text) > 200:
                link_text = link_text[:200] + "..."
            story.append(Paragraph(link_text, normal_style))
    
    story.append(Spacer(1, 15))
    
    # Installation Methods with detailed instructions
    if data.get('installation_methods'):
        story.append(Paragraph("Installation Methods", heading_style))
        story.append(Paragraph("The following installation methods are available for this tool:", normal_style))
        for method in data['installation_methods']:
            method_name = method.get('method', 'Unknown')
            command = method.get('command', '')
            description = method.get('description', '')
            
            method_text = f"For {method_name} installation, use the following command: {command}"
            if len(method_text) > 300:
                method_text = method_text[:300] + "..."
            story.append(Paragraph(method_text, normal_style))
            
            if description:
                desc_text = f"Additional details: {description}"
                if len(desc_text) > 400:
                    desc_text = desc_text[:400] + "..."
                story.append(Paragraph(desc_text, normal_style))
    
    story.append(Spacer(1, 15))
    
    # Dependencies in prose format
    if data.get('dependencies'):
        story.append(Paragraph("Dependencies", heading_style))
        story.append(Paragraph("This tool has the following dependencies:", normal_style))
        for version, deps in data['dependencies'].items():
            deps_text = f"For version {version}, the required dependencies are: {', '.join(deps)}"
            if len(deps_text) > 400:
                deps_text = deps_text[:400] + "..."
            story.append(Paragraph(deps_text, normal_style))
    
    story.append(Spacer(1, 15))
    
    # Category-specific fields in prose
    if data.get('category_fields'):
        story.append(Paragraph("Additional Information", heading_style))
        story.append(Paragraph("The following additional information is available for this tool:", normal_style))
        for key, value in data['category_fields'].items():
            if value:
                field_name = key.replace('_', ' ').title()
                field_text = f"{field_name}: {value}"
                if len(field_text) > 400:
                    field_text = field_text[:400] + "..."
                story.append(Paragraph(field_text, normal_style))
    
    # Build PDF
    doc.build(story)


def get_available_formats() -> list:
    """Get list of available export formats."""
    formats = ['json', 'yaml']
    
    if DOCX_AVAILABLE:
        formats.append('docx')
    
    if PDF_AVAILABLE:
        formats.append('pdf')
    
    return formats 