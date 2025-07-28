"""
Unified Report Processor for MetadataFetcher.

This module provides a single, configurable report generator that can handle
all tool categories with consistent formatting and styling.
"""

import json
import os
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class UnifiedReportGenerator:
    """Unified report generator for all tool categories"""
    
    # Category configurations
    CATEGORY_CONFIGS = {
        'AI_ML': {
            'name': 'AI/ML Development Tools',
            'primary_color': '1F4E79',  # Deep blue
            'secondary_color': '5B9BD5'  # Light blue
        },
        'DATA_SCIENCE': {
            'name': 'Data Science and Analytics Tools',
            'primary_color': '2E7D32',  # Dark green
            'secondary_color': '66BB6A'  # Light green
        },
        'CREATIVE_MEDIA': {
            'name': 'Creative and Media Tools',
            'primary_color': '8E24AA',  # Purple
            'secondary_color': 'BA68C8'  # Light purple
        },
        'DEVELOPER_TOOLS': {
            'name': 'Developer Tools',
            'primary_color': 'D84315',  # Orange
            'secondary_color': 'FF8A65'  # Light orange
        },
        'LLM_TOOLS': {
            'name': 'Large Language Models (LLM) Tools',
            'primary_color': '6A1B9A',  # Deep purple
            'secondary_color': 'AB47BC'  # Light purple
        },
        'GENERIC': {
            'name': 'General Tools',
            'primary_color': '424242',  # Gray
            'secondary_color': '757575'  # Light gray
        }
    }
    
    def __init__(self, category: str = 'GENERIC', output_dir: str = 'outputs'):
        """
        Initialize the unified report generator.
        
        Args:
            category: Tool category (AI_ML, DATA_SCIENCE, etc.)
            output_dir: Base output directory
        """
        self.category = category.upper()
        self.config = self.CATEGORY_CONFIGS.get(self.category, self.CATEGORY_CONFIGS['GENERIC'])
        self.output_dir = output_dir
        self.primary_color = self.config['primary_color']
        self.secondary_color = self.config['secondary_color']
        
    def set_cell_background(self, cell, color: str) -> None:
        """Set cell background color"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
        
    def apply_header_style(self, cell) -> None:
        """Apply header styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
        self.set_cell_background(cell, self.primary_color)
        
    def apply_cell_style(self, cell) -> None:
        """Apply cell styling"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.paragraph_format.space_before = 0
            paragraph.paragraph_format.space_after = 0
            
    def add_section_heading(self, doc, number: int, text: str):
        """Add section heading"""
        heading = doc.add_paragraph()
        run = heading.add_run(f"{number}. {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.line_spacing = 1
        heading.paragraph_format.space_before = 0
        heading.paragraph_format.space_after = 0
        return heading
        
    def add_info_table(self, doc, data: List[tuple]):
        """Create information table"""
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        
        # Set table to use full width within margins
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate column widths
        for row in table.rows:
            row.cells[0].width = Cm(5.5)  # Label column
            row.cells[1].width = Cm(12.96)  # Content column
            
        for i, (key, value) in enumerate(data):
            row_cells = table.rows[i].cells
            
            # Key cell (left column)
            row_cells[0].text = str(key)
            self.apply_header_style(row_cells[0])
            
            # Value cell (right column)
            row_cells[1].text = str(value) if value not in [None, '', [], {}] else 'N/A'
            self.apply_cell_style(row_cells[1])
            
        return table
        
    def add_installation_section(self, doc, installation_data: Dict[str, Any]) -> None:
        """Add installation section"""
        if not installation_data or not isinstance(installation_data, dict):
            doc.add_paragraph('N/A')
            return
            
        # Add installation methods
        if 'methods' in installation_data and installation_data['methods']:
            for method in installation_data['methods']:
                if isinstance(method, dict):
                    method_text = method.get('command', 'N/A')
                    description = method.get('description', '')
                    if description:
                        method_text += f" - {description}"
                    doc.add_paragraph(method_text)
        else:
            doc.add_paragraph('N/A')
            
    def add_dependencies_section(self, doc, dependencies: List[Dict[str, Any]]) -> None:
        """Add dependencies section"""
        if not dependencies:
            doc.add_paragraph('N/A')
            return
            
        for dep in dependencies:
            if isinstance(dep, dict):
                name = dep.get('name', 'Unknown')
                version = dep.get('version', '')
                dep_type = dep.get('type', 'required')
                
                dep_text = f"{name}"
                if version:
                    dep_text += f" ({version})"
                if dep_type != 'required':
                    dep_text += f" [{dep_type}]"
                    
                doc.add_paragraph(dep_text)
        else:
            doc.add_paragraph('N/A')
            
    def add_links_section(self, doc, links: List[Dict[str, Any]]) -> None:
        """Add links section"""
        if not links:
            doc.add_paragraph('N/A')
            return
            
        for link in links:
            if isinstance(link, dict):
                link_type = link.get('type', 'Unknown')
                url = link.get('url', '')
                description = link.get('description', '')
                
                link_text = f"{link_type}: {url}"
                if description:
                    link_text += f" - {description}"
                    
                doc.add_paragraph(link_text)
        else:
            doc.add_paragraph('N/A')
            
    def generate_report(self, metadata: Dict[str, Any], output_path: str) -> None:
        """
        Generate a comprehensive report from metadata.
        
        Args:
            metadata: Tool metadata dictionary
            output_path: Output file path
        """
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = f"{metadata.get('display_name', metadata.get('name', 'Tool'))} Report"
        doc.core_properties.author = "MetadataFetcher"
        
        # Add title
        title = doc.add_heading(f"{metadata.get('display_name', metadata.get('name', 'Tool'))} Report", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add category subtitle
        subtitle = doc.add_paragraph(f"Category: {self.config['name']}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add basic information
        self.add_section_heading(doc, 1, "Basic Information")
        
        basic_info = [
            ("Name", metadata.get('display_name', metadata.get('name', 'N/A'))),
            ("Description", metadata.get('description', 'N/A')),
            ("Version", metadata.get('version', 'N/A')),
            ("Latest Version", metadata.get('latest_version', 'N/A')),
            ("Category", metadata.get('category', 'N/A')),
            ("License", metadata.get('license', 'N/A')),
            ("Author", metadata.get('author', 'N/A')),
            ("Maintainer", metadata.get('maintainer', 'N/A')),
            ("Source Priority", metadata.get('source_priority', 'N/A'))
        ]
        
        self.add_info_table(doc, basic_info)
        
        # Add links section
        self.add_section_heading(doc, 2, "Links")
        links_data = [
            ("Homepage", metadata.get('homepage', 'N/A')),
            ("Documentation", metadata.get('documentation', 'N/A')),
            ("Repository", metadata.get('repository', 'N/A'))
        ]
        self.add_info_table(doc, links_data)
        
        # Add installation section
        self.add_section_heading(doc, 3, "Installation Methods")
        self.add_installation_section(doc, metadata.get('installation_methods', {}))
        
        # Add dependencies section
        self.add_section_heading(doc, 4, "Dependencies")
        self.add_dependencies_section(doc, metadata.get('dependencies', []))
        
        # Add additional links section
        self.add_section_heading(doc, 5, "Additional Links")
        self.add_links_section(doc, metadata.get('links', []))
        
        # Save document
        doc.save(output_path)
        
    def generate_report_from_json(self, json_path: str, output_path: str) -> None:
        """
        Generate report from JSON file.
        
        Args:
            json_path: Path to JSON metadata file
            output_path: Output file path
        """
        try:
            with open(json_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            self.generate_report(metadata, output_path)
        except Exception as e:
            raise Exception(f"Error generating report: {e}")


def gen_report(metadata: Dict[str, Any], category: str, output_path: str) -> None:
    """Generate a unified report for any tool category."""
    generator = UnifiedReportGenerator(category)
    generator.generate_report(metadata, output_path)


def gen_report_from_json(json_path: str, category: str, output_path: str) -> None:
    """Generate a unified report from JSON file."""
    generator = UnifiedReportGenerator(category)
    generator.generate_report_from_json(json_path, output_path)


# Legacy function names for backward compatibility
def export_ai_ml_tool(json_path: str) -> None:
    """Legacy function for AI/ML tools"""
    output_path = json_path.replace('.json', '.docx')
    gen_report_from_json(json_path, 'AI_ML', output_path)


def export_creative_media_tool(json_path: str) -> None:
    """Legacy function for Creative Media tools"""
    output_path = json_path.replace('.json', '.docx')
    gen_report_from_json(json_path, 'CREATIVE_MEDIA', output_path)


def export_data_science_tool(json_path: str) -> None:
    """Legacy function for Data Science tools"""
    output_path = json_path.replace('.json', '.docx')
    gen_report_from_json(json_path, 'DATA_SCIENCE', output_path)


def export_developer_tool(json_path: str) -> None:
    """Legacy function for Developer tools"""
    output_path = json_path.replace('.json', '.docx')
    gen_report_from_json(json_path, 'DEVELOPER_TOOLS', output_path)


def export_llm_tool(json_path: str) -> None:
    """Legacy function for LLM tools"""
    output_path = json_path.replace('.json', '.docx')
    gen_report_from_json(json_path, 'LLM_TOOLS', output_path) 