"""
Advanced output processor for UnifiedMetadata.
Exports metadata to DOCX, PDF, and TXT formats with professional formatting.
"""

import json
import yaml
from typing import Union
from .schema import UnifiedMetadata
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def load_schema():
    """Load the schema configuration."""
    try:
        with open('schema.yaml', 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    except FileNotFoundError:
        # Fallback schema if file not found
        return {
            'basic_info': ['tool_name', 'overview_description', 'key_features', 'use_cases', 'supported_platforms'],
            'technical_specs': ['dependencies', 'system_requirements', 'performance_considerations', 'supported_languages', 'supported_file_formats'],
            'setup_info': ['installation_setup', 'configuration_guide', 'quick_start_tutorial'],
            'documentation': ['official_documentation', 'tutorials_examples', 'community_tutorials', 'api_reference'],
            'community': ['community_support', 'forums_channels', 'ecosystem_packages'],
            'legal_versioning': ['licensing', 'latest_version', 'release_date', 'version_history'],
            'references': ['official_website', 'github_repository', 'documentation_links', 'additional_resources'],
            'category_fields': {},
            'field_descriptions': {},
            'field_types': {}
        }


def get_all_fields():
    """Get all fields from schema in the correct order as defined in schema.yaml."""
    schema = load_schema()
    
    # Get all fields from organized sections
    all_fields = []
    
    # Add fields from each section in order
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    for section in sections:
        section_fields = schema.get(section, [])
        all_fields.extend(section_fields)
    
    # Add category-specific fields
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        for field in fields:
            if field not in all_fields:
                all_fields.append(field)
    
    return all_fields


def get_field_display_name(field_name):
    """Get a proper display name for a field."""
    field_mapping = {
        'tool_name': 'Tool Name',
        'overview_description': 'Overview Description',
        'key_features': 'Key Features',
        'use_cases': 'Use Cases',
        'supported_platforms': 'Supported Platforms',
        'dependencies': 'Dependencies',
        'system_requirements': 'System Requirements',
        'performance_considerations': 'Performance Considerations',
        'supported_languages': 'Supported Languages',
        'supported_file_formats': 'Supported File Formats',
        'hardware_requirements': 'Hardware Requirements',
        'network_requirements': 'Network Requirements',
        'installation_setup': 'Installation Setup',
        'configuration_guide': 'Configuration Guide',
        'quick_start_tutorial': 'Quick Start Tutorial',
        'environment_setup': 'Environment Setup',
        'dependency_management': 'Dependency Management',
        'installation_commands': 'Installation Commands',
        'setup_steps': 'Setup Steps',
        'verification_commands': 'Verification Commands',
        'official_documentation': 'Official Documentation',
        'tutorials_examples': 'Tutorials and Examples',
        'community_tutorials': 'Community Tutorials',
        'api_reference': 'API Reference',
        'video_tutorials': 'Video Tutorials',
        'sample_projects': 'Sample Projects',
        'community_support': 'Community Support',
        'forums_channels': 'Forums and Channels',
        'ecosystem_packages': 'Ecosystem Packages',
        'support_channels': 'Support Channels',
        'user_groups': 'User Groups',
        'licensing': 'Licensing',
        'latest_version': 'Latest Version',
        'release_date': 'Release Date',
        'version_history': 'Version History',
        'update_policy': 'Update Policy',
        'end_of_life': 'End of Life',
        'official_website': 'Official Website',
        'github_repository': 'GitHub Repository',
        'documentation_links': 'Documentation Links',
        'additional_resources': 'Additional Resources',
        'download_links': 'Download Links',
        'demo_links': 'Demo Links',
        # AI/ML specific fields
        'model_types_supported': 'Supported Model Types',
        'training_capabilities': 'Training Capabilities',
        'inference_features': 'Inference Features',
        'hardware_acceleration': 'Hardware Acceleration',
        'ml_framework_integration': 'ML Framework Integration',
        'model_deployment_options': 'Model Deployment Options',
        'example_notebooks': 'Example Notebooks',
        'pre_trained_models': 'Pre-trained Models',
        'model_optimization': 'Model Optimization',
        'data_preprocessing_tools': 'Data Preprocessing Tools',
        # Data Science specific fields
        'data_formats_supported': 'Supported Data Formats',
        'visualization_capabilities': 'Visualization Capabilities',
        'statistical_analysis': 'Statistical Analysis',
        'data_processing_features': 'Data Processing Features',
        'ml_library_integration': 'ML Library Integration',
        'data_export_formats': 'Data Export Formats',
        'interactive_features': 'Interactive Features',
        'data_import_tools': 'Data Import Tools',
        'reporting_features': 'Reporting Features',
        'data_quality_tools': 'Data Quality Tools',
        # Creative/Media specific fields
        'media_formats_supported': 'Supported Media Formats',
        'rendering_capabilities': 'Rendering Capabilities',
        'plugin_ecosystem': 'Plugin Ecosystem',
        'user_interface': 'User Interface',
        'collaboration_features': 'Collaboration Features',
        'export_formats': 'Export Formats',
        'performance_optimization': 'Performance Optimization',
        'asset_management': 'Asset Management',
        'workflow_automation': 'Workflow Automation',
        'real_time_features': 'Real-time Features',
        # Developer Tools specific fields
        'language_support': 'Language Support',
        'ide_integration': 'IDE Integration',
        'version_control_integration': 'Version Control Integration',
        'debugging_tools': 'Debugging Tools',
        'code_analysis': 'Code Analysis',
        'build_deployment': 'Build and Deployment',
        'extension_ecosystem': 'Extension Ecosystem',
        'testing_frameworks': 'Testing Frameworks',
        'code_generation': 'Code Generation',
        'refactoring_tools': 'Refactoring Tools',
        # LLM Tools specific fields
        'supported_model_types': 'Supported Model Types',
        'prompt_engineering_tools': 'Prompt Engineering Tools',
        'fine_tuning_capabilities': 'Fine-tuning Capabilities',
        'api_sdk_availability': 'API/SDK Availability',
        'model_management': 'Model Management',
        'deployment_options': 'Deployment Options',
        'integration_frameworks': 'Integration Frameworks',
        'evaluation_metrics': 'Evaluation Metrics',
        'safety_features': 'Safety Features',
        'cost_optimization': 'Cost Optimization'
    }
    
    return field_mapping.get(field_name, field_name.replace('_', ' ').title())


def get_section_display_name(section_name):
    """Get a proper display name for a section."""
    section_mapping = {
        'basic_info': 'Basic Information',
        'technical_specs': 'Technical Specifications',
        'setup_info': 'Setup and Installation',
        'documentation': 'Documentation and Learning',
        'community': 'Community and Support',
        'legal_versioning': 'Legal and Versioning',
        'references': 'References and Links'
    }
    
    return section_mapping.get(section_name, section_name.replace('_', ' ').title())


def format_field_value(value):
    """Format a field value for display, handling lists and long text appropriately."""
    if isinstance(value, list):
        if len(value) == 1:
            return str(value[0])
        elif len(value) <= 3:
            # Short lists: comma-separated
            return ", ".join(str(item) for item in value)
        else:
            # Long lists: numbered format but more compact
            formatted_items = []
            for i, item in enumerate(value, 1):
                formatted_items.append(f"{i}. {str(item)}")
            return "\n".join(formatted_items)
    elif isinstance(value, dict):
        # Handle dictionaries by converting to key-value pairs
        if not value:
            return ""
        formatted_items = []
        for key, val in value.items():
            formatted_items.append(f"{key}: {val}")
        return "\n".join(formatted_items)
    else:
        # For long text, break into paragraphs but be more conservative
        text = str(value)
        if len(text) > 300:  # Higher threshold for breaking
            # Split by common delimiters and format as numbered list
            if ". " in text:
                sentences = text.split(". ")
                formatted_items = []
                for i, sentence in enumerate(sentences, 1):
                    if sentence.strip():
                        formatted_items.append(f"{i}. {sentence.strip()}")
                return "\n".join(formatted_items)
            elif " - " in text:
                items = text.split(" - ")
                formatted_items = []
                for i, item in enumerate(items, 1):
                    if item.strip():
                        formatted_items.append(f"{i}. {item.strip()}")
                return "\n".join(formatted_items)
            else:
                # Just return the text as is for very long content
                return text
        else:
            return text


def save_to_docx(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a DOCX file with professional formatting.
    Follows user preferences: numbered headers, 11pt Times New Roman, narrow margins, proper spacing.
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create document
    doc = Document()
    
    # Set narrow margins (1.27 cm)
    section = doc.sections[0]
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    
    # Set up custom styles
    # Title style - 11pt Times New Roman, black, bold
    title_style = doc.styles['Title']
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(11)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    # Heading 1 style - 11pt Times New Roman, black, bold, no spacing after
    heading_style = doc.styles['Heading 1']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(11)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)  # Black
    heading_style.paragraph_format.space_after = Pt(0)  # No spacing after
    heading_style.paragraph_format.space_before = Pt(0)  # No extra spacing
    
    # Normal style for body text - 11pt Times New Roman, readable spacing
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)  # Minimal spacing
    normal_style.paragraph_format.space_before = Pt(0)  # No extra spacing
    normal_style.paragraph_format.line_spacing = 1.2  # Readable spacing
    
    # Add title
    tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
    title = doc.add_heading(f"{tool_name} - Metadata Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process fields by sections with proper headers and logical ordering
    schema = load_schema()
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    section_number = 1
    
    # Process each section in order
    for section in sections:
        section_fields = schema.get(section, [])
        section_data = data.get(section, {})
        
        # Check if section has any meaningful content
        has_content = False
        for field in section_fields:
            value = section_data.get(field)
            if value and value != "" and value != [] and value != {}:
                has_content = True
                break
        
        if has_content:
            # Add section header
            section_title = get_section_display_name(section)
            doc.add_heading(f"{section_number}. {section_title}", level=1)
            section_number += 1
            
            # Process fields within this section
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    # Add field header (no extra spacing)
                    field_title = get_field_display_name(field)
                    doc.add_heading(f"  {field_title}", level=2)
                    
                    # Add field content in prose format (immediately after header)
                    if isinstance(value, list):
                        # Convert list to prose
                        if len(value) == 1:
                            doc.add_paragraph(str(value[0]))
                        else:
                            prose_text = f"{value[0]}"
                            for item in value[1:-1]:
                                prose_text += f", {item}"
                            if len(value) > 1:
                                prose_text += f", and {value[-1]}"
                            doc.add_paragraph(prose_text)
                    else:
                        formatted_value = format_field_value(value)
                        if formatted_value and formatted_value.strip():
                            doc.add_paragraph(formatted_value)
    
    # Process category-specific fields (at root level)
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        for field in fields:
            value = data.get(field)
            if value and value != "" and value != [] and value != {}:
                # Add field header (no extra spacing)
                field_title = get_field_display_name(field)
                doc.add_heading(f"{section_number}. {field_title}", level=1)
                
                # Add field content in prose format (immediately after header)
                if isinstance(value, list):
                    # Convert list to prose
                    if len(value) == 1:
                        doc.add_paragraph(str(value[0]))
                    else:
                        prose_text = f"{value[0]}"
                        for item in value[1:-1]:
                            prose_text += f", {item}"
                        if len(value) > 1:
                            prose_text += f", and {value[-1]}"
                        doc.add_paragraph(prose_text)
                else:
                    formatted_value = format_field_value(value)
                    if formatted_value and formatted_value.strip():
                        doc.add_paragraph(formatted_value)
                
                section_number += 1
    
    # Save document
    doc.save(file_path)


def save_to_pdf(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a PDF file with professional formatting using tables for all content.
    Follows user preferences: numbered headers, 11pt Times New Roman, narrow margins, proper spacing.
    """
    if not PDF_AVAILABLE:
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
    
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    # Create PDF document with narrow margins (1.27 cm)
    doc = SimpleDocTemplate(
        file_path, 
        pagesize=A4,
        leftMargin=1.27*cm,
        rightMargin=1.27*cm,
        topMargin=1.27*cm,
        bottomMargin=1.27*cm
    )
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles matching user preferences
    # Title style - 11pt Times New Roman, bold, centered
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=12,
        alignment=1,  # Center
        textColor=colors.black
    )
    
    # Normal style for body text - 11pt Times New Roman, readable spacing
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        spaceAfter=6,
        spaceBefore=0,
        leading=14,  # Readable spacing
        textColor=colors.black,
        keepWithNext=False,
        keepWithPrevious=False
    )
    
    # Header style - 11pt Times New Roman, bold, no spacing after
    header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading1'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=0,
        spaceBefore=0,
        textColor=colors.black,
        keepWithNext=True,
        keepWithPrevious=False
    )
    
    # Table header style
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontName='Times-Bold',
        fontSize=11,
        spaceAfter=0,
        spaceBefore=0,
        textColor=colors.white,
        alignment=0  # Left
    )
    
    # Table cell style with better multi-line support and proper line breaking
    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        spaceAfter=0,
        spaceBefore=0,
        textColor=colors.black,
        alignment=0,  # Left
        leading=13,  # Tighter line spacing for better fit
        keepWithNext=False,
        keepWithPrevious=False,
        wordWrap='CJK',  # Better word wrapping
        splitLongWords=True,  # Break long words if needed
    )
    
    # Title
    tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
    story.append(Paragraph(f"{tool_name} - Metadata Report", title_style))
    
    # Process fields by sections with proper headers and logical ordering
    schema = load_schema()
    sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
    
    section_number = 1
    
    # Process each section in order
    for section in sections:
        section_fields = schema.get(section, [])
        section_data = data.get(section, {})
        
        # Check if section has any meaningful content
        has_content = False
        for field in section_fields:
            value = section_data.get(field)
            if value and value != "" and value != [] and value != {}:
                has_content = True
                break
        
        if has_content:
            # Add section header
            section_title = get_section_display_name(section)
            story.append(Paragraph(f"{section_number}. {section_title}", header_style))
            story.append(Spacer(1, 4))  # Add small space after section header
            section_number += 1
            
            # Create table for all fields in this section
            table_data = []
            table_data.append([Paragraph("Field", table_header_style), Paragraph("Details", table_header_style)])
            
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    field_title = get_field_display_name(field)
                    
                    # Format value based on type
                    if isinstance(value, list):
                        if len(value) == 1:
                            details = str(value[0])
                        else:
                            details = ", ".join(str(item) for item in value)
                    else:
                        details = format_field_value(value)
                    
                    # Special formatting for installation_setup with step-by-step guide
                    if field == 'installation_setup' and details:
                        details = format_installation_guide(details)
                    
                    table_data.append([
                        Paragraph(field_title, table_cell_style),
                        Paragraph(details, table_cell_style)
                    ])
            
            if len(table_data) > 1:  # Has data beyond header
                # Calculate available width: A4 width (21cm) - 2*1.27cm margins = 18.46cm
                available_width = 18.46 * cm
                # Use 30% for field names, 70% for details
                field_width = available_width * 0.3
                details_width = available_width * 0.7
                
                table = Table(table_data, colWidths=[field_width, details_width])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F81BD')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 4),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))  # Reduced spacing between sections
    
    # Process category-specific fields (at root level) with table format for better organization
    category_fields = schema.get('category_fields', {})
    for category, fields in category_fields.items():
        # Check if any category fields have content
        category_data = []
        for field in fields:
            value = data.get(field)
            if value and value != "" and value != [] and value != {}:
                category_data.append((field, value))
        
        if category_data:
            story.append(Paragraph(f"{section_number}. {category} Features", header_style))
            story.append(Spacer(1, 4))  # Add small space after section header
            section_number += 1
            
            # Create table for category-specific features
            table_data = []
            table_data.append([Paragraph("Feature", table_header_style), Paragraph("Details", table_header_style)])
            
            for field, value in category_data:
                field_title = get_field_display_name(field)
                details = format_field_value(value)
                
                table_data.append([
                    Paragraph(field_title, table_cell_style),
                    Paragraph(details, table_cell_style)
                ])
            
            if len(table_data) > 1:  # Has data beyond header
                # Calculate available width: A4 width (21cm) - 2*1.27cm margins = 18.46cm
                available_width = 18.46 * cm
                # Use 30% for field names, 70% for details
                field_width = available_width * 0.3
                details_width = available_width * 0.7
                
                table = Table(table_data, colWidths=[field_width, details_width])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4F81BD')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 11),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 4),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))  # Reduced spacing between sections
    
    # Build PDF
    doc.build(story)


def format_installation_guide(installation_text: str) -> str:
    """
    Format installation text into a detailed step-by-step guide with commands.
    """
    if not installation_text:
        return ""
    
    # Split installation methods
    methods = installation_text.split('. ')
    formatted_guide = []
    
    for i, method in enumerate(methods, 1):
        if not method.strip():
            continue
            
        # Parse method name and description
        if ':' in method:
            parts = method.split(':', 1)
            method_name = parts[0].strip()
            description = parts[1].strip() if len(parts) > 1 else ""
        elif ' - ' in method:
            parts = method.split(' - ', 1)
            method_name = parts[0].strip()
            description = parts[1].strip() if len(parts) > 1 else ""
        else:
            method_name = method.strip()
            description = ""
        
        # Generate step-by-step guide with commands
        step_guide = f"Step {i}: {method_name}\n"
        
        if description:
            step_guide += f"Description: {description}\n"
        
        # Add specific commands based on method type
        if "pip" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  pip install [package-name]\n"
            step_guide += "  pip install --upgrade [package-name]\n"
        elif "docker" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  docker pull [image-name]\n"
            step_guide += "  docker run [image-name]\n"
        elif "git" in method_name.lower() or "github" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  git clone [repository-url]\n"
            step_guide += "  cd [project-directory]\n"
            step_guide += "  pip install -r requirements.txt\n"
        elif "conda" in method_name.lower() or "anaconda" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  conda install [package-name]\n"
            step_guide += "  conda create -n [env-name] python=3.x\n"
        elif "npm" in method_name.lower() or "node" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  npm install [package-name]\n"
            step_guide += "  npm install -g [package-name]\n"
        elif "brew" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  brew install [package-name]\n"
            step_guide += "  brew update && brew upgrade\n"
        elif "apt" in method_name.lower() or "ubuntu" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  sudo apt update\n"
            step_guide += "  sudo apt install [package-name]\n"
        elif "yum" in method_name.lower() or "centos" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  sudo yum update\n"
            step_guide += "  sudo yum install [package-name]\n"
        elif "chocolatey" in method_name.lower():
            step_guide += "Commands:\n"
            step_guide += "  choco install [package-name]\n"
            step_guide += "  choco upgrade [package-name]\n"
        elif "official" in method_name.lower() or "installer" in method_name.lower():
            step_guide += "Steps:\n"
            step_guide += "  1. Download installer from official website\n"
            step_guide += "  2. Run the installer as administrator\n"
            step_guide += "  3. Follow the installation wizard\n"
            step_guide += "  4. Add to PATH if required\n"
        else:
            step_guide += "Steps:\n"
            step_guide += "  1. Follow the official documentation\n"
            step_guide += "  2. Ensure prerequisites are met\n"
            step_guide += "  3. Run installation commands\n"
            step_guide += "  4. Verify installation\n"
        
        formatted_guide.append(step_guide)
    
    return "\n\n".join(formatted_guide)


def save_to_txt(metadata: Union[UnifiedMetadata, dict], file_path: str) -> None:
    """
    Save UnifiedMetadata to a TXT file with detailed, structured formatting.
    Follows the format shown in sample_report.txt with proper section headers and comprehensive content.
    """
    if isinstance(metadata, UnifiedMetadata):
        data = metadata.to_dict()
    else:
        data = metadata
    
    with open(file_path, 'w', encoding='utf-8') as f:
        # Title
        tool_name = data.get('basic_info', {}).get('tool_name', data.get('tool_name', data.get('display_name', data.get('name', 'Tool'))))
        f.write(f"{tool_name} - Metadata Report\n\n")
        
        # Process fields by sections with proper headers and logical ordering
        schema = load_schema()
        sections = ['basic_info', 'technical_specs', 'setup_info', 'documentation', 'community', 'legal_versioning', 'references']
        
        section_number = 1
        
        # Process each section in order
        for section in sections:
            section_fields = schema.get(section, [])
            section_data = data.get(section, {})
            
            # Check if section has any meaningful content
            has_content = False
            for field in section_fields:
                value = section_data.get(field)
                if value and value != "" and value != [] and value != {}:
                    has_content = True
                    break
            
            if has_content:
                # Add section header
                section_title = get_section_display_name(section)
                f.write(f"{section_number}. {section_title}\n")
                section_number += 1
                
                # Process fields within this section
                for field in section_fields:
                    value = section_data.get(field)
                    if value and value != "" and value != [] and value != {}:
                        # Add field header (no extra spacing)
                        field_title = get_field_display_name(field)
                        f.write(f"  {field_title}\n")
                        
                        # Add field content in prose format (immediately after header)
                        if isinstance(value, list):
                            # Convert list to prose
                            if len(value) == 1:
                                f.write(f"    {value[0]}\n")
                            else:
                                prose_text = f"{value[0]}"
                                for item in value[1:-1]:
                                    prose_text += f", {item}"
                                if len(value) > 1:
                                    prose_text += f", and {value[-1]}"
                                f.write(f"    {prose_text}\n")
                        else:
                            formatted_value = format_field_value(value)
                            if formatted_value and formatted_value.strip():
                                f.write(f"    {formatted_value}\n")
                        f.write("\n")
        
        # Process category-specific fields (at root level)
        category_fields = schema.get('category_fields', {})
        for category, fields in category_fields.items():
            for field in fields:
                value = data.get(field)
                if value and value != "" and value != [] and value != {}:
                    # Add field header (no extra spacing)
                    field_title = get_field_display_name(field)
                    f.write(f"{section_number}. {field_title}\n")
                    
                    # Add field content in prose format (immediately after header)
                    if isinstance(value, list):
                        # Convert list to prose
                        if len(value) == 1:
                            f.write(f"  {value[0]}\n")
                        else:
                            prose_text = f"{value[0]}"
                            for item in value[1:-1]:
                                prose_text += f", {item}"
                            if len(value) > 1:
                                prose_text += f", and {value[-1]}"
                            f.write(f"  {prose_text}\n")
                    else:
                        formatted_value = format_field_value(value)
                        if formatted_value and formatted_value.strip():
                            f.write(f"  {formatted_value}\n")
                    
                    section_number += 1
                    f.write("\n")


def get_available_formats() -> list:
    """Get list of available export formats."""
    formats = ['json', 'yaml', 'txt']
    
    if DOCX_AVAILABLE:
        formats.append('docx')
    
    if PDF_AVAILABLE:
        formats.append('pdf')
    
    return formats 